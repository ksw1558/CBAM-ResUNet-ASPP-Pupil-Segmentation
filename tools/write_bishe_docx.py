from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT_PATH = Path.home() / "Desktop" / "毕设.docx"
BACKUP_PATH = Path.home() / "Desktop" / "毕设_原文件备份.docx"


def set_run_font(run, size=12, bold=False, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 13, bold=True, font="黑体")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=12, font="宋体")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("• " + text)
    set_run_font(run, size=12, font="宋体")
    return p


def set_cell(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(str(text))
    set_run_font(run, size=10.5, bold=bold, font="宋体")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            set_cell(cells[i], item)
    doc.add_paragraph()


def prepare_doc():
    if OUT_PATH.exists() and not BACKUP_PATH.exists():
        BACKUP_PATH.write_bytes(OUT_PATH.read_bytes())

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.7)
    sec.right_margin = Cm(2.7)
    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(12)
    return doc


def main():
    doc = prepare_doc()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("毕业设计知识补充：基于深度学习的瞳孔分割、中心定位与直径测量")
    set_run_font(run, size=18, bold=True, font="黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("结合 Pupil_Segmentation_Project 项目内容整理")
    set_run_font(run, size=12, font="宋体")

    add_heading(doc, "一、课题研究背景与意义")
    add_body(doc, "本课题围绕瞳孔区域的精确分割、瞳孔中心坐标提取以及瞳孔直径计算展开，属于计算机视觉、医学图像处理和人机交互方向的交叉研究。瞳孔位置和瞳孔直径是眼动分析中的重要参数，可用于注视点估计、视觉注意分析、疲劳检测、VR/AR交互以及认知负荷评估等应用场景。")
    add_body(doc, "在VR眼动场景中，眼部图像通常受到红外光照、眼睑遮挡、睫毛干扰、反光、低对比度和运动模糊等因素影响。传统阈值分割或边缘检测方法在复杂场景下稳定性不足，因此本毕业设计采用深度学习语义分割模型提取瞳孔区域，并结合椭圆拟合实现几何参数估计。")

    add_heading(doc, "二、项目整体技术路线")
    add_body(doc, "本项目的核心流程可以概括为：输入眼部图像或视频帧，利用改进的UNet类网络进行瞳孔二分类分割，得到瞳孔mask；随后对mask进行连通域筛选和形态学处理，提取最大瞳孔区域；再利用cv2.fitEllipse进行椭圆拟合，获得瞳孔中心、长轴、短轴和角度；最后根据长短轴计算瞳孔直径，并生成可视化图、统计表和论文实验图。")
    add_bullet(doc, "图像分割阶段：模型输出每个像素属于瞳孔区域的概率。")
    add_bullet(doc, "后处理阶段：二值化、去噪、最大连通域筛选和孔洞填充。")
    add_bullet(doc, "几何估计阶段：椭圆拟合得到中心坐标(cx, cy)、major_axis、minor_axis。")
    add_bullet(doc, "实验分析阶段：计算mIoU、Dice、Recall、中心误差、置信度和瞳孔直径统计指标。")

    add_heading(doc, "三、数据集与实验数据说明")
    add_body(doc, "项目中原先将主要训练数据称为VOC数据集，但从研究内容上看，其本质是LPW瞳孔数据，因此论文中使用“LPW数据集”表述更加严谨。OpenEDS数据集用于跨数据集测试，能够反映模型在不同眼动采集条件下的泛化能力。项目后续还尝试了TEyeD-like、CASIA等补充数据或展示样本，用于增强论文实验的完整性。")
    add_table(doc, ["数据/目录", "用途", "论文中建议表述"], [
        ["LPW / VOCdevkit", "主要训练与验证数据", "LPW数据集，用于模型训练与主实验评估"],
        ["resources/videos", "视频输入", "用于瞳孔中心与直径的连续帧分析"],
        ["teacher_report_assets", "论文图表与汇报素材", "用于整理对比实验、消融实验和可视化结果"],
    ])

    add_heading(doc, "四、最终模型结构说明：CBAM-ResUNet-ASPP")
    add_body(doc, "本项目最终选定实验11模型作为毕业设计主模型，模型名称为CBAM-ResUNet-ASPP。该模型以UNet编码器—解码器结构为基础，引入残差连接、CBAM注意力机制以及ASPP多尺度上下文模块，目的是提升瞳孔边界分割精度和复杂场景下的鲁棒性。")
    add_table(doc, ["模块", "作用", "对瞳孔分割的意义"], [
        ["UNet结构", "编码器提取语义特征，解码器恢复空间分辨率", "适合医学/眼部图像这类像素级分割任务"],
        ["Residual Block", "缓解深层网络梯度退化，增强特征表达", "提高模型训练稳定性和边界细节保留能力"],
        ["CBAM注意力", "从通道和空间两个维度强调关键区域", "让模型更加关注瞳孔区域，抑制眼睑、睫毛和反光干扰"],
        ["ASPP模块", "通过不同膨胀率卷积提取多尺度上下文", "适应不同大小、不同形态的瞳孔区域"],
    ])
    add_body(doc, "实验11最终模型在LPW数据集上的mIoU约为98.16%，作为毕业设计最终模型较为合适。论文中应避免只强调单一mIoU数值，而应结合Dice、Recall、预测mask图、中心定位结果和瞳孔直径分析共同证明模型有效性。")

    add_heading(doc, "五、损失函数与模型优化思路")
    add_body(doc, "指导教师曾建议从loss入手提升mIoU，这一思路是合理的。瞳孔分割任务不仅要求区域分割准确，还要求边界平滑、中心稳定，因此损失函数可以综合考虑区域重叠、边界约束和几何中心约束。项目中涉及Basic Loss、Edge Loss、Centroid Loss等消融设置，能够体现模型优化的针对性。")
    add_table(doc, ["损失/约束", "主要目的", "论文说明角度"], [
        ["Basic Loss", "保证像素级分类正确", "作为基础分割损失"],
        ["Dice Loss / IoU Loss", "提升前景区域重叠程度", "适合前景面积较小的瞳孔分割"],
        ["Edge Loss", "增强边界约束", "改善瞳孔边缘粗糙问题"],
        ["Centroid Loss", "约束预测区域中心稳定性", "与最终瞳孔中心定位目标一致"],
    ])

    add_heading(doc, "六、评价指标说明")
    add_body(doc, "论文实验部分应清楚解释每个指标的含义。mIoU反映预测mask与真实mask的平均交并比，是语义分割任务的核心指标；Dice反映两个区域的重叠程度；Recall反映真实瞳孔区域被检出的比例。对于最终应用目标，还应补充中心误差和瞳孔直径统计。")
    add_body(doc, "mIoU = TP / (TP + FP + FN)，其中TP表示正确预测为瞳孔的像素，FP表示错误预测为瞳孔的背景像素，FN表示漏检的瞳孔像素。Dice = 2TP / (2TP + FP + FN)。Recall = TP / (TP + FN)。")
    add_table(doc, ["指标", "含义", "适用实验"], [
        ["mIoU", "预测区域与真实区域交并比", "分割主实验、对比实验、消融实验"],
        ["Dice", "区域重叠程度", "分割质量评估"],
        ["Recall", "瞳孔区域召回能力", "漏检情况分析"],
        ["Center Error", "预测中心与真实中心的欧氏距离", "瞳孔中心定位实验"],
        ["Pupil Diameter", "由拟合椭圆长短轴平均得到", "瞳孔直径变化分析"],
    ])

    add_heading(doc, "七、瞳孔中心定位方法")
    add_body(doc, "瞳孔中心定位并不是直接由网络回归得到，而是基于分割mask进行几何计算。具体做法是先将模型输出的mask二值化，然后提取最大连通域作为瞳孔区域。若该区域轮廓点数量不少于5个，则使用OpenCV中的cv2.fitEllipse拟合椭圆，拟合椭圆的中心即为瞳孔中心(cx, cy)。")
    add_body(doc, "为了提升鲁棒性，项目中还设计了置信度评价指标，综合考虑最大连通域面积占比、椭圆拟合质量和圆度。置信度越高，说明分割区域越集中、几何形态越稳定，中心定位结果越可信。")
    add_body(doc, "中心误差的计算公式为：Center Error = sqrt((cx_pred - cx_gt)^2 + (cy_pred - cy_gt)^2)。该指标可直接反映模型在瞳孔定位任务中的实际应用价值。")

    add_heading(doc, "八、瞳孔直径计算方法")
    add_body(doc, "瞳孔直径基于椭圆拟合结果计算。由于真实瞳孔在图像中常因视角、眼球姿态和成像畸变呈现椭圆形，因此直接使用圆形假设并不稳定。项目采用椭圆长轴major_axis和短轴minor_axis的平均值作为像素级瞳孔直径：pupil_diameter = (major_axis + minor_axis) / 2。")
    add_body(doc, "在视频分析中，逐帧计算瞳孔直径可以得到随时间变化的直径序列。为了降低帧间噪声，项目使用移动平均进行平滑处理，例如rolling mean窗口设置为5。平滑后的曲线更适合论文展示，可以反映瞳孔直径的整体变化趋势。")

    add_heading(doc, "九、对比实验与消融实验设计")
    add_body(doc, "毕业设计中不仅需要报告最终模型结果，还应通过对比实验和消融实验说明模型优势。对比实验用于证明CBAM-ResUNet-ASPP相较于Base U-Net、CA-UNet、CBAM-UNet、TransUNet、VM-UNet、DeepLabV3+、PSPNet、ResUNet等模型的综合表现；消融实验用于证明各模块和损失函数的贡献。")
    add_table(doc, ["实验类型", "目的", "建议展示内容"], [
        ["模型对比实验", "证明最终模型总体性能较优", "mIoU、Dice、Recall表格和综合对比图"],
        ["跨数据集实验", "验证泛化能力", "其他瞳孔数据上的指标与预测图"],
        ["消融实验", "验证模块有效性", "w/o CBAM、w/o ASPP、w/o Edge Loss等设置"],
        ["定性可视化", "展示mask边界和困难样本效果", "Input、Ground Truth、Prediction、Error Map"],
    ])

    add_heading(doc, "十、论文图表整理建议")
    add_body(doc, "论文图表应避免PPT风格，尽量采用白底、Times New Roman字体、统一字号、简洁配色和清晰图例。对于分割结果图，建议使用Input、Ground Truth、Prediction、Error Map四列结构；对于消融图，建议使用横向条形图展示ΔmIoU；对于瞳孔中心定位图，建议展示GT椭圆、预测椭圆、中心点和误差箭头；对于瞳孔直径分析图，建议展示原始曲线、平滑曲线、均值和标准差。")
    add_table(doc, ["图表名称", "对应文件/目录", "论文用途"], [
        ["综合性能对比表", "teacher_report_assets", "展示不同模型在LPW上的指标"],
        ["消融实验表和图", "teacher_report_assets/ablation_plan", "证明CBAM、ASPP、损失函数等模块有效"],
        ["中心定位可视化", "11_CBAM-ResUNet-ASPP/pupil_params_results", "展示几何解释性和中心误差"],
        ["直径变化曲线", "11_CBAM-ResUNet-ASPP/video_diameter_analysis", "展示视频中瞳孔直径随时间变化"],
    ])

    add_heading(doc, "十一、项目主要目录说明")
    add_table(doc, ["目录/文件", "作用说明"], [
        ["11_CBAM-ResUNet-ASPP", "最终模型实验目录，包含训练、预测、中心提取和直径分析代码"],
        ["nets/cbam_res_unet_exp11.py", "实验11最终模型结构定义"],
        ["logs/final_exp11_miou98_16_epoch070.pth", "最终模型权重文件"],
        ["resources/videos", "用于视频瞳孔中心和直径分析的输入视频"],
        ["video_diameter_results_full", "逐帧瞳孔直径CSV结果"],
        ["video_diameter_analysis", "瞳孔直径统计表和论文图"],
        ["teacher_report_assets", "给老师检查和论文使用的图表素材"],
    ])

    add_heading(doc, "十二、论文创新点表述建议")
    add_body(doc, "本毕业设计的创新点不应只写“使用UNet进行瞳孔分割”，而应突出针对瞳孔任务的结构改进和应用闭环。可以概括为：构建CBAM-ResUNet-ASPP瞳孔分割模型；结合注意力机制与多尺度上下文增强复杂场景下的瞳孔区域提取；引入边界和中心约束思想提高分割与定位稳定性；在分割结果基础上完成瞳孔中心和直径的几何估计，实现从mask预测到眼动参数提取的完整流程。")
    add_bullet(doc, "创新点1：面向瞳孔分割任务设计CBAM-ResUNet-ASPP结构，增强局部边界和全局上下文表达。")
    add_bullet(doc, "创新点2：结合消融实验分析CBAM、ASPP及损失函数组件对分割性能的影响。")
    add_bullet(doc, "创新点3：基于分割mask进行椭圆拟合，实现瞳孔中心坐标和瞳孔直径的自动提取。")
    add_bullet(doc, "创新点4：面向视频场景生成瞳孔直径时间序列和统计图，为VR眼动分析提供可解释参数。")

    add_heading(doc, "十三、小结")
    add_body(doc, "综上，本项目从语义分割模型设计、模型训练与评估、跨数据集验证、消融实验、定性可视化、瞳孔中心定位和瞳孔直径分析等方面形成了较完整的毕业设计工作链条。论文撰写时应围绕“高精度瞳孔分割是基础，中心坐标和直径提取是最终应用目标”这一主线展开，避免只停留在mIoU指标描述上。")

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")
    print(f"Backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
